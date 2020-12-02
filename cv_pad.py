from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

#profile pic
document.add_picture(
    'pad.jpg',
    width = Inches(2.0)
    )

#nombre/fono/mail
speak('¿Cuál es tu nombre?')
nombre = input('¿Cuál es tu nombre? ')
telefono = input('¿Cuál es tu número de teléfono? ')
email = input('¿Cuál es tu correo electrónico? ')

document.add_paragraph( nombre + ' | ' + telefono + ' | ' + email)

#perfil
document.add_heading('Perfil')
document.add_paragraph( input('Cuantanos un poco sobre tí: ')
)

#experiencia
document.add_heading('Experiencia Laboral')
p = document.add_paragraph()

company = input('Nombre de la empresa: ')
start_date = input('Fecha de inicio: ')
end_date = input('Fecha de fin: ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + ' / ' + end_date + '\n').italic = True
p.add_run(input('Cuentanos un poco sobre tu experiencia en ' + company + ': ' + '\n'))

#mas experiencia?
while True:
    mas_exp = input('Agregar más experiencia? S/N: ')
    if mas_exp.lower() == 's':
        p = document.add_paragraph()

        company = input('Nombre de la empresa: ')
        start_date = input('Fecha de inicio: ')
        end_date = input('Fecha de fin: ')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + ' / ' + end_date + '\n').italic = True
        p.add_run(input('Cuentanos un poco sobre tu experiencia en ' + company + ': ' + '\n'))
    else:
        break

#habilidades
document.add_heading('Habilidades', level = 1)

#habilidades tecnicas
document.add_heading('Técnicas', level = 2)
document.add_paragraph(
    input('Habilidades Técnicas. Ingresa una: '),
    style = 'List Bullet'
    )

while True:
    mas_hab_tec = input('Quieres agregar alguna otra habilidad técnica?: S/N ')
    if mas_hab_tec.lower() == 's':
        document.add_paragraph(
            input('Habilidades Técnicas. Ingresa una: '),
            style = 'List Bullet'
            )
    else:
        break

#habilidades personales
document.add_heading('Personales', level = 2)
document.add_paragraph(
    input('Habilidades Personales. Ingresa una: '),
    style = 'List Bullet'
    )

while True:
    mas_hab_pers = input('Quieres agregar alguna otra habilidad Personal?: S/N ')
    if mas_hab_pers.lower() == 's':
        document.add_paragraph(
            input('Habilidades Personales. Ingresa una: '),
            style = 'List Bullet'
            )
    else:
        break



document.save('cv.docx')